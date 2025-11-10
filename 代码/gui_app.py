import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import json
import os
import sys
import time
import asyncio
import requests
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from 主干.生成文案 import process_texts
from 支线.教育.json转word import json_to_word

# 导入图片爬取功能
IMAGE_CRAWLING_AVAILABLE = False
key_school_finded = None
get_school_image = None
set_login_callback = None
is_image_relevant_to_school = None

try:
    from 支线.教育.爬取图片 import key_school_finded, get_school_image, set_login_callback
    # 尝试导入其他函数
    try:
        from 支线.教育.爬取图片 import is_image_relevant_to_school
    except ImportError:
        is_image_relevant_to_school = None
    IMAGE_CRAWLING_AVAILABLE = True
except ImportError:
    print("图片爬取功能不可用，请安装所需依赖")

class XiaohongshuApp:
    def __init__(self, root):
        self.root = root
        self.root.title("小红书文案生成器")
        self.root.geometry("800x600")
        
        # 设置当前工作目录为脚本所在目录
        os.chdir(os.path.dirname(os.path.abspath(__file__)))
        
        # 控制标志
        self.is_running = False
        self.is_waiting_for_login = False
        
        # 文案类型变量
        self.content_type = tk.StringVar(value="学校介绍")
        
        # 线程引用
        self.worker_thread = None
        
        # 创建界面
        self.create_widgets()
        
        # 设置登录回调
        if IMAGE_CRAWLING_AVAILABLE and set_login_callback:
            set_login_callback(self.on_need_login)
        
    def create_widgets(self):
        # 创建标签页
        tab_control = ttk.Notebook(self.root)
        
        # 文案生成标签页
        self.text_generation_tab = ttk.Frame(tab_control)
        tab_control.add(self.text_generation_tab, text="文案生成")
        
        # 图片爬取标签页
        self.image_crawling_tab = ttk.Frame(tab_control)
        tab_control.add(self.image_crawling_tab, text="图片爬取")
        
        # 使用说明标签页
        self.instructions_tab = ttk.Frame(tab_control)
        tab_control.add(self.instructions_tab, text="使用说明")
        
        tab_control.pack(expand=1, fill="both")
        
        # 文案生成标签页内容
        self.create_text_generation_tab()
        
        # 图片爬取标签页内容
        self.create_image_crawling_tab()
        
        # 使用说明标签页内容
        self.create_instructions_tab()
        
    def create_instructions_tab(self):
        # 创建文本框显示使用说明
        instructions_frame = ttk.Frame(self.instructions_tab)
        instructions_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建滚动文本框
        text_widget = scrolledtext.ScrolledText(instructions_frame, wrap=tk.WORD, padx=10, pady=10)
        text_widget.pack(fill=tk.BOTH, expand=True)
        
        # 读取使用说明文件内容

        with open("使用说明.md", "r", encoding="utf-8") as f:
            content = f.read()
            text_widget.insert(tk.END, content)
        
        # 禁止编辑
        text_widget.config(state=tk.DISABLED)
        
    def create_text_generation_tab(self):
        # 文案类型选择
        type_frame = ttk.LabelFrame(self.text_generation_tab, text="文案类型")
        type_frame.grid(row=0, column=0, columnspan=3, padx=10, pady=5, sticky="ew")
        
        # 添加红星标记表示必填项
        ttk.Label(type_frame, text="*", foreground="red").grid(row=0, column=0, padx=(5,0), pady=5, sticky="w")
        ttk.Radiobutton(type_frame, text="学校介绍", variable=self.content_type, value="学校介绍").grid(row=0, column=1, padx=5, pady=5)
        ttk.Radiobutton(type_frame, text="升学指导", variable=self.content_type, value="升学指导").grid(row=0, column=2, padx=5, pady=5)
        
        # 文件选择
        file_frame = ttk.LabelFrame(self.text_generation_tab, text="文件选择")
        file_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=5, sticky="ew")
        
        # 添加红星标记表示必填项
        ttk.Label(file_frame, text="*", foreground="red").grid(row=0, column=0, padx=(5,0), pady=5, sticky="w")
        ttk.Label(file_frame, text="输入文件:").grid(row=0, column=0, padx=(20,0), pady=5, sticky="w")
        self.input_file_entry = ttk.Entry(file_frame, width=50)
        self.input_file_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(file_frame, text="浏览", command=self.browse_input_file).grid(row=0, column=2, padx=5, pady=5)
        
        ttk.Label(file_frame, text="*", foreground="red").grid(row=1, column=0, padx=(5,0), pady=5, sticky="w")
        ttk.Label(file_frame, text="输出目录:").grid(row=1, column=0, padx=(20,0), pady=5, sticky="w")
        self.output_dir_entry = ttk.Entry(file_frame, width=50)
        self.output_dir_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(file_frame, text="浏览", command=self.browse_output_dir).grid(row=1, column=2, padx=5, pady=5)
        
        ttk.Label(file_frame, text="输出文件名:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.output_filename_entry = ttk.Entry(file_frame, width=50)
        self.output_filename_entry.grid(row=2, column=1, padx=5, pady=5, sticky="ew")
        ttk.Label(file_frame, text=".json").grid(row=2, column=2, padx=5, pady=5, sticky="w")
        
        file_frame.columnconfigure(1, weight=1)
        
        # 控制按钮
        control_frame = ttk.Frame(self.text_generation_tab)
        control_frame.grid(row=2, column=0, columnspan=3, padx=10, pady=5, sticky="ew")
        
        self.control_button = ttk.Button(control_frame, text="开始生成文案", command=self.toggle_generation)
        self.control_button.pack(side=tk.LEFT, padx=5)
        
        # 进度条
        self.progress = ttk.Progressbar(control_frame, mode='indeterminate')
        self.progress.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # 日志显示
        log_frame = ttk.LabelFrame(self.text_generation_tab, text="处理日志")
        log_frame.grid(row=3, column=0, columnspan=3, padx=10, pady=5, sticky="nsew")
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=20)
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.text_generation_tab.columnconfigure(0, weight=1)
        self.text_generation_tab.rowconfigure(3, weight=1)
        
    def create_image_crawling_tab(self):
        # Word文档选择
        ttk.Label(self.image_crawling_tab, text="Word文档:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.word_file_var = tk.StringVar()
        ttk.Entry(self.image_crawling_tab, textvariable=self.word_file_var, width=50).grid(row=0, column=1, sticky="ew", pady=5)
        ttk.Button(self.image_crawling_tab, text="浏览", command=self.browse_word_file).grid(row=0, column=2, padx=(5,0), pady=5)
        
        # 输出目录选择
        ttk.Label(self.image_crawling_tab, text="输出目录:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.image_output_dir_var = tk.StringVar()
        ttk.Entry(self.image_crawling_tab, textvariable=self.image_output_dir_var, width=50).grid(row=1, column=1, sticky="ew", pady=5)
        ttk.Button(self.image_crawling_tab, text="浏览", command=self.browse_image_output_dir).grid(row=1, column=2, padx=(5,0), pady=5)
        
        # 控制按钮
        self.image_control_button = ttk.Button(self.image_crawling_tab, text="开始爬取", command=self.toggle_image_crawling)
        self.image_control_button.grid(row=2, column=0, columnspan=3, pady=20)
        
        # 进度条
        self.image_progress = ttk.Progressbar(self.image_crawling_tab, mode='indeterminate')
        self.image_progress.grid(row=3, column=0, columnspan=3, sticky="ew", pady=10)
        
        # 日志显示区域
        ttk.Label(self.image_crawling_tab, text="运行日志:").grid(row=4, column=0, sticky=tk.W, pady=(10, 5))
        self.image_log_text = scrolledtext.ScrolledText(self.image_crawling_tab, height=15)
        self.image_log_text.grid(row=5, column=0, columnspan=3, sticky="ewns", pady=5)
        
        # 配置网格权重
        self.image_crawling_tab.columnconfigure(1, weight=1)
        self.image_crawling_tab.rowconfigure(5, weight=1)
        
    def browse_input_file(self):
        filename = filedialog.askopenfilename(
            title="选择原始文案文件",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if filename:
            self.input_file_entry.delete(0, tk.END)
            self.input_file_entry.insert(0, filename)
            
    def browse_output_dir(self):
        directory = filedialog.askdirectory(title="选择输出目录")
        if directory:
            self.output_dir_entry.delete(0, tk.END)
            self.output_dir_entry.insert(0, directory)
            
    def browse_word_file(self):
        filename = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.word_file_var.set(filename)
            
    def browse_image_output_dir(self):
        directory = filedialog.askdirectory(title="选择输出目录")
        if directory:
            self.image_output_dir_var.set(directory)
            
    def log_message(self, *args):
        """日志消息处理函数，支持可变参数"""
        if len(args) == 1:
            # 单个消息参数
            message = args[0]
        elif len(args) == 3:
            # 三个参数：类型、索引、内容
            msg_type, index, content = args
            if msg_type == "original":
                message = f"[{index}] 原始文案:\n{content}\n"
            elif msg_type == "rewrite":
                message = f"[{index}] 需要重写文案\n"
            elif msg_type == "result":
                message = f"[{index}] 最终文案:\n{content}\n"
            else:
                message = f"[{index}] {content}\n"
        else:
            # 其他情况，将所有参数转换为字符串
            message = " ".join(str(arg) for arg in args)
            
        # 在主线程中更新UI
        self.root.after(0, self._update_log_display, message)
        
    def _update_log_display(self, message):
        """在主线程中更新日志显示"""
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.log_text.update()
        
    def log_image_message(self, message):
        self.image_log_text.insert(tk.END, message + "\n")
        self.image_log_text.see(tk.END)
        self.root.update_idletasks()
        
    def toggle_generation(self):
        """切换生成状态：开始/停止"""
        if not self.is_running:
            # 开始生成
            self.start_generation()
        else:
            # 停止生成
            self.stop_generation()
            
    def start_generation(self):
        input_file = self.input_file_entry.get()
        output_dir = self.output_dir_entry.get()
        content_type = self.content_type.get()
        output_filename = self.output_filename_entry.get().strip()
        
        # 如果用户没有输入文件名，使用默认文件名
        if not output_filename:
            output_filename = "finally_school_texts"
        
        # 确保文件名以.json结尾
        if not output_filename.endswith('.json'):
            output_filename += '.json'
        
        if not input_file or not output_dir:
            messagebox.showerror("错误", "请选择输入文件和输出目录")
            return
            
        if not os.path.exists(input_file):
            messagebox.showerror("错误", "输入文件不存在")
            return
            
        if not os.path.exists(output_dir):
            messagebox.showerror("错误", "输出目录不存在")
            return
            
        # 设置输出文件路径
        output_file = os.path.join(output_dir, output_filename)
        
        self.is_running = True
        self.control_button.config(text="停止生成")
        self.progress.start()
        
        # 在新线程中运行生成过程
        def run_generation():
            try:
                # 传递文案类型参数
                process_texts(input_file, output_file, content_type=content_type, callback=self.log_message)
                # 生成完成后自动转为Word文档
                if self.is_running:
                    json_to_word(output_file)  # json_to_word函数只需要一个参数
                    word_file = output_file.replace(".json", ".docx")  # 生成的Word文件名
                    self.log_message(f"Word文档已生成: {word_file}")
                    self.root.after(0, lambda: messagebox.showinfo("完成", "文案生成和Word转换已完成!"))
            except Exception as e:
                self.log_message(f"处理过程中出现错误: {str(e)}")
                self.root.after(0, lambda: messagebox.showerror("错误", f"处理过程中出现错误: {str(e)}"))
            finally:
                self.is_running = False
                self.root.after(0, self.on_generation_finished)
                
        self.worker_thread = threading.Thread(target=run_generation, daemon=True)
        self.worker_thread.start()
        
    def stop_generation(self):
        """停止生成"""
        self.is_running = False
        self.control_button.config(text="开始生成")
        self.progress.stop()
        self.log_message("=== 已停止生成 ===")
        
    def on_generation_finished(self):
        self.control_button.config(text="开始生成")
        self.progress.stop()
        
    def toggle_image_crawling(self):
        """切换图片爬取状态：开始/停止"""
        if not hasattr(self, 'is_image_crawling'):
            self.is_image_crawling = False
            
        if not self.is_image_crawling:
            # 开始爬取
            self.start_image_crawling()
        else:
            # 停止爬取
            self.stop_image_crawling()
            
    def start_image_crawling(self):
        word_file = self.word_file_var.get()
        output_dir = self.image_output_dir_var.get()
        
        if not word_file or not output_dir:
            messagebox.showerror("错误", "请选择Word文档和输出目录")
            return
            
        if not os.path.exists(word_file):
            messagebox.showerror("错误", "Word文档不存在")
            return
            
        # 生成输出文件路径
        word_filename = os.path.basename(word_file)
        name_without_ext = os.path.splitext(word_filename)[0]
        output_file = os.path.join(output_dir, f"{name_without_ext}_img.docx")
        
        # 设置运行状态
        self.is_image_crawling = True
        self.image_control_button.config(text="停止爬取")
        self.image_progress.start()
        
        # 在新线程中执行图片爬取任务
        self.image_worker_thread = threading.Thread(target=self.crawl_images, args=(word_file, output_file))
        self.image_worker_thread.daemon = True
        self.image_worker_thread.start()
        
    def stop_image_crawling(self):
        """停止图片爬取"""
        self.is_image_crawling = False
        self.image_control_button.config(text="开始爬取")
        self.image_progress.stop()
        self.log_image_message("=== 已停止图片爬取 ===")
        
    def on_need_login(self):
        """当需要登录时调用此函数"""
        self.is_waiting_for_login = True
        self.log_image_message("检测到需要登录小红书账号")
        self.log_image_message("请在打开的浏览器中扫码登录，登录完成后点击'登录完成'按钮")
        
        # 在主线程中显示登录提示对话框
        self.root.after(0, self.show_login_dialog)
        
        # 等待用户点击登录完成按钮
        while self.is_waiting_for_login:
            time.sleep(0.1)
            self.root.update()
            
    def show_login_dialog(self):
        """显示登录提示对话框"""
        # 创建顶层窗口
        login_window = tk.Toplevel(self.root)
        login_window.title("需要登录")
        login_window.geometry("300x150")
        login_window.transient(self.root)
        login_window.grab_set()  # 模态对话框
        
        # 居中显示
        login_window.geometry("+%d+%d" % (login_window.winfo_screenwidth()/2-150, login_window.winfo_screenheight()/2-75))
        
        # 提示信息
        tk.Label(login_window, text="请在打开的浏览器中扫码登录小红书账号", 
                wraplength=280, justify=tk.CENTER).pack(pady=20)
        
        # 登录完成按钮
        tk.Button(login_window, text="登录完成", 
                 command=lambda: self.on_login_completed(login_window)).pack(pady=10)
        
    def on_login_completed(self, login_window):
        """用户点击登录完成按钮"""
        self.is_waiting_for_login = False
        login_window.destroy()
        self.log_image_message("登录完成，继续图片爬取...")
        
    def crawl_images(self, word_file, output_file):
        if not IMAGE_CRAWLING_AVAILABLE or key_school_finded is None or get_school_image is None:
            self.log_image_message("图片爬取功能不可用")
            return
            
        try:
            self.log_image_message("开始爬取图片...")
            
            # 加载Word文档
            doc = Document(word_file)
            self.log_image_message("成功加载Word文档")
            
            # 创建输出文档（如果不存在）
            if not os.path.exists(output_file):
                output_doc = Document()
                # 设置页面边距
                sections_output = output_doc.sections
                for section in sections_output:
                    section.top_margin = Cm(2)
                    section.bottom_margin = Cm(2)
                    section.left_margin = Cm(2)
                    section.right_margin = Cm(2)
                # 保存初始空文档
                output_doc.save(output_file)
                self.log_image_message(f"创建新的Word文档: {output_file}")
            else:
                self.log_image_message(f"使用现有的Word文档: {output_file}")
            
            # 获取所有章节
            sections = []
            current_section = {}
            
            # 遍历所有段落，根据Heading 2样式分割
            for paragraph in doc.paragraphs:
                # 检查是否停止
                if not self.is_image_crawling:
                    self.log_image_message("=== 已停止 ===")
                    return
                    
                # 检查段落样式是否为Heading 2
                if paragraph.style and paragraph.style.name == 'Heading 2':
                    # 如果已经有当前章节，保存它
                    if current_section:
                        sections.append(current_section)
                    # 开始新章节
                    current_section = {
                        'title': paragraph.text,
                        'content': ''
                    }
                else:
                    # 如果是普通段落，添加到当前章节内容中
                    if current_section and paragraph.text and isinstance(paragraph.text, str) and paragraph.text.strip():
                        if current_section['content']:
                            current_section['content'] += '\n' + paragraph.text
                        else:
                            current_section['content'] = paragraph.text
            
            # 添加最后一个章节
            if current_section:
                sections.append(current_section)
            
            self.log_image_message(f"找到 {len(sections)} 个章节")
            
            # 检查已处理的章节数量，实现断点续跑功能
            processed_sections = 0
            if os.path.exists(output_file):
                try:
                    output_doc_check = Document(output_file)
                    # 统计Heading 2的数量来确定已处理的章节数
                    heading_count = 0
                    for paragraph in output_doc_check.paragraphs:
                        if paragraph.style and paragraph.style.name == 'Heading 2':
                            heading_count += 1
                    processed_sections = heading_count
                    self.log_image_message(f"已处理 {processed_sections} 个章节")
                except Exception as e:
                    self.log_image_message(f"检查已处理章节时出错: {e}")
                    processed_sections = 0
            
            # 确保images文件夹存在
            if not os.path.exists('images'):
                os.makedirs('images')
                self.log_image_message("创建images文件夹")
            
            # 处理每个章节（从已处理的章节开始）
            for i, section in enumerate(sections[processed_sections:], processed_sections):
                # 检查是否停止
                if not self.is_image_crawling:
                    self.log_image_message("=== 已停止 ===")
                    return
                    
                title = section['title']
                content = section['content']
                text = title + '\n' + content
                self.log_image_message(f"处理第 {i+1}/{len(sections)} 章节: {title}")
                
                # 通过key_school_finded找到学校名称
                try:
                    school_name = key_school_finded(text)
                    if school_name:
                        self.log_image_message(f"找到学校: {school_name}")
                        
                        # 调用图片爬取功能
                        self.log_image_message("正在获取学校图片...")
                        try:
                            image_urls = get_school_image(school_name)
                            self.log_image_message(f"获取到 {len(image_urls) if image_urls else 0} 张图片")
                            
                            if image_urls:
                                self.log_image_message("图片获取成功")
                                # 下载并处理图片
                                self.log_image_message("开始下载图片...")
                                downloaded_images = self.download_images(image_urls)
                                self.log_image_message(f"下载了 {len(downloaded_images)} 张图片")
                                
                                # AI审核图片
                                self.log_image_message("开始AI审核图片...")
                                filtered_images = self.filter_images_with_ai(downloaded_images)
                                self.log_image_message(f"AI审核后保留 {len(filtered_images)} 张图片")
                                
                                # 将文案和图片添加到Word文档
                                self.log_image_message("将文案和图片添加到Word文档...")
                                self.add_content_to_word(output_file, title, content, filtered_images)
                            else:
                                self.log_image_message("未获取到图片")
                                # 即使没有图片，也要将文案添加到Word文档
                                self.add_content_to_word(output_file, title, content, [])
                        except Exception as e:
                            self.log_image_message(f"获取图片时出错: {str(e)}")
                            # 即使获取图片失败，也要将文案添加到Word文档
                            self.add_content_to_word(output_file, title, content, [])
                    else:
                        self.log_image_message("未找到学校名称")
                        # 即使未找到学校，也要将文案添加到Word文档
                        self.add_content_to_word(output_file, title, content, [])
                except Exception as e:
                    self.log_image_message(f"查找学校名称时出错: {str(e)}")
                    # 即使出错，也要将文案添加到Word文档
                    self.add_content_to_word(output_file, title, content, [])
                
                # 在每次处理完一个章节后检查是否停止
                if not self.is_image_crawling:
                    self.log_image_message("=== 已停止 ===")
                    return
                    
                time.sleep(1)  # 避免请求过于频繁
            
            self.log_image_message("图片爬取完成!")
            self.log_image_message(f"生成的文件: {output_file}")
            
            messagebox.showinfo("完成", "图片爬取已完成!")
        except Exception as e:
            self.log_image_message(f"处理过程中出现错误: {str(e)}")
            messagebox.showerror("错误", f"处理过程中出现错误: {str(e)}")
        finally:
            self.image_progress.stop()
            self.image_control_button.config(text="开始爬取")
            self.is_image_crawling = False
            
    def download_images(self, image_urls):
        """下载图片并返回本地路径列表"""
        downloaded_images = []
        for i, url in enumerate(image_urls[:10]):  # 限制最多10张图片
            # 检查是否停止
            if not self.is_image_crawling:
                self.log_image_message("=== 已停止 ===")
                return downloaded_images
                
            try:
                self.log_image_message(f"下载图片 {i+1}/{len(image_urls)}: {url}")
                response = requests.get(url, timeout=10)
                if response.status_code == 200:
                    # 生成文件名
                    filename = f"image_{i+1}.jpg"
                    file_path = os.path.join('images', filename)
                    
                    # 保存图片
                    with open(file_path, 'wb') as f:
                        f.write(response.content)
                    downloaded_images.append(file_path)
                    self.log_image_message(f"已下载图片: {filename}")
                else:
                    self.log_image_message(f"下载图片失败，状态码: {response.status_code}")
            except Exception as e:
                self.log_image_message(f"下载图片 {url} 时出错: {e}")
        return downloaded_images
        
    def filter_images_with_ai(self, image_paths):
        """使用AI审核图片，返回符合要求的图片路径列表"""
        if not is_image_relevant_to_school:
            self.log_image_message("AI图片审核功能不可用，返回所有图片")
            return image_paths
            
        filtered_images = []
        for image_path in image_paths:
            # 检查是否停止
            if not self.is_image_crawling:
                self.log_image_message("=== 已停止 ===")
                return filtered_images
                
            try:
                time.sleep(2)  # 避免请求过于频繁
                if is_image_relevant_to_school(image_path):
                    filtered_images.append(image_path)
                    self.log_image_message(f"✅ 图片通过AI审核: {os.path.basename(image_path)}")
                else:
                    self.log_image_message(f"❌ 图片未通过AI审核: {os.path.basename(image_path)}")
                    # 删除不符合的图片
                    try:
                        os.remove(image_path)
                        self.log_image_message(f"已删除不符合主题的图片: {os.path.basename(image_path)}")
                    except Exception as e:
                        self.log_image_message(f"删除图片 {os.path.basename(image_path)} 失败: {e}")
            except Exception as e:
                self.log_image_message(f"AI审核图片 {os.path.basename(image_path)} 时出错: {e}")
                # 出错时保守处理：保留图片
                filtered_images.append(image_path)
        return filtered_images
        
    def add_content_to_word(self, output_file, title, content, image_paths):
        """将文案和图片添加到Word文档并立即保存"""
        try:
            # 打开现有的Word文档
            doc = Document(output_file)
            
            # 插入筛选后的图片
            if image_paths:
                for image_path in image_paths:
                    try:
                        if os.path.exists(image_path):
                            self.log_image_message(f"准备插入图片: {os.path.basename(image_path)}")
                            # 检查文件大小
                            file_size = os.path.getsize(image_path)
                            if file_size == 0:
                                self.log_image_message(f"图片文件为空，跳过: {os.path.basename(image_path)}")
                                continue
                            
                            # 插入图片
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
                            run = paragraph.add_run()
                            # 使用较小的图片尺寸，避免内存问题
                            run.add_picture(image_path, width=Inches(4))
                            doc.add_paragraph()  # 添加空行
                            self.log_image_message(f"插入图片成功: {os.path.basename(image_path)}")
                        else:
                            self.log_image_message(f"图片文件不存在: {os.path.basename(image_path)}")
                    except Exception as insert_error:
                        self.log_image_message(f"插入图片 {os.path.basename(image_path)} 时出错: {insert_error}")
                        # 尝试使用PIL处理图片
                        try:
                            from PIL import Image
                            from io import BytesIO
                            # 打开图片并转换为RGB模式
                            with Image.open(image_path) as img:
                                # 转换为RGB模式（如果是RGBA或其他模式）
                                if img.mode in ('RGBA', 'LA', 'P'):
                                    img = img.convert('RGB')
                                # 调整图片大小以减少内存使用
                                max_size = (1200, 1200)  # 最大尺寸
                                img.thumbnail(max_size, Image.Resampling.LANCZOS)
                                # 保存为JPEG格式到内存
                                img_buffer = BytesIO()
                                img.save(img_buffer, format='JPEG', quality=85)
                                img_buffer.seek(0)
                                # 插入图片
                                paragraph = doc.add_paragraph()
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = paragraph.add_run()
                                run.add_picture(img_buffer, width=Inches(4))
                                doc.add_paragraph()  # 添加空行
                                self.log_image_message(f"通过PIL插入图片成功: {os.path.basename(image_path)}")
                        except Exception as pil_error:
                            self.log_image_message(f"使用PIL插入图片 {os.path.basename(image_path)} 时出错: {pil_error}")
                            # 如果PIL也失败，尝试插入错误信息
                            error_para = doc.add_paragraph()
                            error_para.add_run(f"[图片插入失败: {os.path.basename(image_path)}]").bold = True
            
            # 插入标题和内容
            if title:
                heading = doc.add_heading(title, level=2)
            if content:
                doc.add_paragraph(content)
            
            # 添加分页符
            doc.add_page_break()
            
            # 立即保存文档
            doc.save(output_file)
            self.log_image_message("Word文档已保存")
            
        except Exception as e:
            self.log_image_message(f"添加内容到Word文档时出错: {e}")

    def run_async_func(self, func, *args):
        """在新线程中运行异步函数"""
        try:
            # 创建新的事件循环
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            # 确保传递的是可等待的对象
            if asyncio.iscoroutinefunction(func):
                coro = func(*args)
            else:
                # 如果不是协程函数，直接调用
                result = func(*args)
                loop.close()
                return result
            result = loop.run_until_complete(coro)
            loop.close()
            return result
        except Exception as e:
            self.log_image_message(f"异步函数执行出错: {str(e)}")
            return []

def main():
    root = tk.Tk()
    app = XiaohongshuApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()