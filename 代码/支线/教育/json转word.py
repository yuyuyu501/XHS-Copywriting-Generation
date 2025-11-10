import json
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def json_to_word(file_path):

    # 创建Word文档
    doc = Document()
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    with open(file_path, "r", encoding="utf-8") as f:
        school_texts = json.load(f)
        for school_text in school_texts:
            # 清理文本内容
            school_text = school_text.replace("\n", "")
            school_text = school_text.replace("：", ":")
            school_text = school_text.replace(" ", "")
            title_part = school_text.split("标题:")[1]
            title = title_part.split("文案:")[0]
            content = "".join(title_part.split("文案:")[1:])

            # 添加文档标题
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(14)
            title_run.bold = True
            title_para.style = 'Heading 2'
            
            # 添加内容
            content_para = doc.add_paragraph()
            content_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            content_run = content_para.add_run(content)
            content_run.font.size = Pt(11)

            # 分页
            doc.add_page_break()

        # 保存Word文档
        word_file = file_path.replace(".json", ".docx")
        doc.save(word_file)
        print(f"成功生成Word文档: {word_file}")

        return True



if __name__ == '__main__':
    json_file = 'D://小红书文案2//文案//教育//生成文案//finally_school_texts.json'
    json_to_word(json_file)