import time
import pandas as pd
from playwright.async_api import async_playwright
import asyncio
from typing import List, Dict, Any
import os
import requests
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from openai import OpenAI
import base64

# 配置日志
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)
logging.getLogger("openai").setLevel(logging.WARNING)
logging.getLogger("httpcore").setLevel(logging.WARNING)
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("urllib3").setLevel(logging.WARNING)
logging.getLogger("requests").setLevel(logging.WARNING)

# 添加全局变量来存储登录状态
need_login_callback = None

def set_login_callback(callback):
    """设置登录回调函数"""
    global need_login_callback
    need_login_callback = callback

# 初始化 OpenAI 客户端（你已提供）
client = OpenAI(
    base_url="https://ark.cn-beijing.volces.com/api/v3",
    api_key="baea195a-fe33-4be1-9c68-9292a4a66a5c",
)
MODEL = "doubao-seed-1-6-vision-250815"

prompt = """# 角色
你是一个图片主题判断助手，能够精准判断图片是否符合探校这个主题。

## 技能
### 技能 1: 判断图片主题
1. 当用户提供图片时，仔细分析图片内容，要求是现实的照片。
2. 防止出现教师招聘、招聘信息等无关内容。
3. 根据图片中的场景、元素、文字、人物活动等方面，判断该图片是否符合探校主题。
4. 给出明确的判断结果，回复示例如下：
   - 判断结果: 符合：符合
              不符合：不符合   
   - 理由: <详细说明判断的依据，如图片中出现了校园建筑、学生在校园的活动等表明符合；若图片内容为自然风光等则表明不符合>

## 限制:
- 只专注于判断图片是否符合探校主题，拒绝回答与该主题判断无关的话题。
- 所输出的内容必须按照给定的格式进行组织，不能偏离框架要求。
- 判断理由需简洁明了，控制在合理字数内。"""

def key_school_finded(text: str) -> str:
    """
    获取学校名称

    参数：
    text: 待处理的文本
    
    返回：
    学校名称
    """
    # logger.debug(f"开始查找文本中的学校名称: {text}")
    
    ## 从guangzhou_all_schools.csv文件中获取学校名称
    # 构造正确的CSV文件路径
    csv_path = os.path.join(os.path.dirname(__file__), 'guangzhou_all_schools.csv')
    if not os.path.exists(csv_path):
        # 尝试其他可能的路径
        csv_path = 'guangzhou_all_schools.csv'
        if not os.path.exists(csv_path):
            csv_path = '../guangzhou_all_schools.csv'
            if not os.path.exists(csv_path):
                csv_path = '../../guangzhou_all_schools.csv'
    schools_df = pd.read_csv(csv_path)
    logger.debug(f"成功读取CSV文件，共有 {len(schools_df)} 行数据")
    
    ## 遍历schools_df中的所有列，查找学校名称
    for column in schools_df.columns:
        logger.debug(f"检查列: {column}")
        for school in schools_df[column]:
            if isinstance(school, str) and school in text:
                logger.debug(f"找到匹配的学校: {school}")
                return school
    ### 如果没有找到匹配的学校，返回空字符串
    logger.debug("未找到匹配的学校")
    return ""


def get_school_image(name: str) -> List:
    """
    根据学校名称获取学校图片

    参数：
    name: 学校名称

    返回：
    学校图片的URL
    """
    logger.debug(f"开始获取学校 {name} 的图片")
    # 使用异步函数获取图片
    result = asyncio.run(_get_school_image_async(name))
    logger.debug(f"获取到 {len(result)} 张图片")
    return result


async def _get_school_image_async(school_name: str) -> List:
    """
    异步版本的学校图片获取函数
    根据学校名称搜索小红书，按点赞数排序，获取第一篇笔记的所有图片URL
    """
    logger.debug(f"异步函数开始执行，搜索学校: {school_name}")
    
    if not school_name:
        logger.debug("学校名称为空，返回空列表")
        return []
    
    async with async_playwright() as p:
        # 启动浏览器（关闭无头模式以便观察）
        logger.debug("启动浏览器（可视化模式）")
        browser = await p.chromium.launch(headless=False)
        # 使用持久化上下文来保存和重用cookies
        user_data_dir = "支线//教育"
        if not os.path.exists(user_data_dir):
            os.makedirs(user_data_dir)
        context = await browser.new_context(storage_state=user_data_dir + "/xhs_cookies.json" if os.path.exists(user_data_dir + "/xhs_cookies.json") else None)
        page = None
        
        try:
            # 构造搜索URL（添加默认按热度排序）
            from urllib.parse import quote
            encoded_name = quote(school_name)
            search_url = f"https://www.xiaohongshu.com/search_result?keyword={encoded_name}"
            logger.debug(f"搜索URL: {search_url}")
            
            # 访问搜索页面
            page = await context.new_page()
            logger.debug("访问搜索页面")
            await page.goto(search_url)
            await page.wait_for_timeout(5000)  # 等待页面加载
            logger.debug("页面加载完成")
            
            # 检查是否需要登录
            try:
                await page.wait_for_selector("text=登录", timeout=5000)
                logger.debug("检测到需要登录")
                
                # 使用回调函数通知GUI需要登录
                if need_login_callback:
                    need_login_callback()
                else:
                    print("检测到需要登录，请扫码后按回车继续...")
                    input()
                
                # 保存cookies
                await context.storage_state(path=user_data_dir + "/xhs_cookies.json")
                logger.debug("已保存最新cookie")
            except Exception:
                logger.debug("已登录，无需扫码")
            
            # 确保按“最多点赞”排序（修复版：确保筛选按钮可点击）
            logger.debug("正在设置排序方式为‘最多点赞’...")
            try:
                # 1. 等待“筛选”按钮出现并可交互
                logger.debug("等待筛选按钮出现...")
                await page.wait_for_selector("div.filter", state="visible", timeout=8000)
                logger.debug("筛选按钮已可见，准备点击...")

                # 可选：滚动到元素位置，避免被遮挡
                await page.eval_on_selector("div.filter", "el => el.scrollIntoViewIfNeeded()")

                # 点击筛选按钮
                await page.click("div.filter", timeout=5000)
                logger.debug("已点击‘筛选’按钮，展开菜单")

                # 2. 等待筛选面板出现
                await page.wait_for_selector("div.filter-panel", state="visible", timeout=5000)
                logger.debug("筛选菜单面板已成功展开")

                # 3. 点击“最多点赞”选项（使用文本 + 可见性）
                await page.wait_for_selector("text=最多点赞", state="visible", timeout=5000)
                await page.click("text=最多点赞")
                logger.debug("已成功点击‘最多点赞’选项")

                # 4. 等待笔记列表重新加载
                await page.wait_for_selector("div.note-item", timeout=8000)
                logger.debug("排序后的内容已成功加载")

            except Exception as e:
                logger.warning(f"无法设置‘最多点赞’排序: {e}")
                # 可在此处添加 fallback：比如直接使用 &sort=popularity_descending 的 URL

            # 滚动页面以加载更多内容
            logger.debug("滚动页面以加载更多内容")
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(3000)  # 作为加载更多内容的缓冲

            # 滚动页面以加载更多内容
            logger.debug("滚动页面以加载更多内容")
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(3000)  # 作为加载更多内容的缓冲
            
            # 滚动页面以加载更多内容
            logger.debug("滚动页面以加载更多内容")
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(3000)
            
            # 查找笔记链接和封面图
            logger.debug("开始查找笔记链接和封面图")
            
            # 定义规范化链接函数
            BASE_URL = "https://www.xiaohongshu.com"
            
            def _normalize_note_link(raw_href: str) -> str | None:
                """将各种笔记链接规范化为可直接打开的笔记详情页链接"""
                if not raw_href:
                    return None

                # 需要保留查询参数（token 等），因此先拆分但同时保留原始串
                original = raw_href
                href_path_only = raw_href
                if "?" in raw_href:
                    href_path_only = raw_href.split("?")[0]

                # 去掉协议主机前缀，统一成 path 处理（仅用于判断类型，不改变 original）
                if href_path_only.startswith("http"):
                    try:
                        from urllib.parse import urlparse
                        parsed = urlparse(href_path_only)
                        href_path_only = parsed.path
                    except Exception:
                        pass

                # /user/profile/<uid>/<noteId> 规范化
                if "/user/profile/" in href_path_only:
                    # 对于用户主页卡片给到的更准确链接，直接保留原始链接（含查询参数）
                    # original 可能是绝对或相对地址
                    if original.startswith("http"):
                        return original
                    # 相对地址补全域名
                    return f"{BASE_URL}{original}"

                # /explore/<noteId> 直接返回
                if "/explore/" in href_path_only:
                    note_id = href_path_only.split("/explore/")[-1].strip("/")
                    if note_id:
                        return f"{BASE_URL}/explore/{note_id}"

                # /discovery/item/<noteId> 规范化
                if "/discovery/item/" in href_path_only:
                    note_id = href_path_only.split("/discovery/item/")[-1].strip("/")
                    if note_id:
                        return f"{BASE_URL}/explore/{note_id}"

                return None
            
            # 尝试多种选择器来获取笔记链接和封面图（按优先级排序）
            link_selectors = [
                "a.cover.mask.ld[href]",     # 最优先：封面卡片链接
                "a[href*='/explore/']",      # explore链接（笔记详情页）
                "a[href*='/discovery/item/']",  # discovery链接（笔记详情页）
                ".note-item a[href]",        # 笔记项链接
                ".feeds-container a[href]",  # feed容器中的链接
                "[class*='note' i] a[href]", # 包含note的链接
                "a[href*='/user/profile/']"  # 用户主页格式的链接（最低优先级）
            ]
            
            image_urls = []
            note_url = None
            
            # 首先尝试直接获取封面图片
            image_selectors = [
                "a.cover.mask.ld img[src]",   # 封面卡片中的图片
                "img[data-xhs-img][src]",     # 带有data-xhs-img属性的图片
                ".note-item img[src]",        # 笔记项中的图片
                ".feeds-container img[src]",  # feed容器中的图片
                "img[src*='xhscdn.com']"      # 小红书CDN图片
            ]
            
            # 尝试直接获取封面图片URL
            MAX_IMAGES = 10
            for selector in image_selectors:
                if len(image_urls) >= MAX_IMAGES:
                    break
                try:
                    logger.debug(f"尝试图片选择器: {selector}")
                    imgs = await page.query_selector_all(selector)
                    logger.debug(f"找到 {len(imgs)} 个图片元素")
                    for img in imgs:
                        if len(image_urls) >= MAX_IMAGES:
                            break
                        src = await img.get_attribute("src")
                        if src and src not in image_urls:
                            image_urls.append(src)
                            logger.debug(f"找到封面图片URL: {src}")
                except Exception as e:
                    logger.error(f"图片选择器 {selector} 处理出错: {e}")
                    continue
            
            await browser.close()
            logger.debug(f"总共获取到 {len(image_urls)} 张图片")
            return image_urls[:10]  # 返回前10张图片
            
        except Exception as e:
            logger.error(f"获取学校图片时出错: {e}")
            # 尝试获取页面截图用于调试
            try:
                if page:
                    await page.screenshot(path="debug_error_page.png", full_page=True)
                    logger.debug("已保存错误页面截图用于调试")
            except Exception as screenshot_error:
                logger.error(f"保存错误截图时出错: {screenshot_error}")
            await browser.close()
            return []


def is_image_relevant_to_school(image_path: str) -> bool:
    """
    使用视觉大模型判断图片是否符合"探校"主题
    返回 True 表示符合，False 表示不符合
    """
    logger.debug(f"正在使用AI判断图片是否符合探校主题: {image_path}")
    
    try:
        # 读取图片并转为 base64
        with open(image_path, "rb") as image_file:
            encoded_image = base64.b64encode(image_file.read()).decode('utf-8')
        
        # 构造消息 - 使用正确的类型
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{encoded_image}"
                        }
                    }
                ]
            }
        ]

        # 调用模型
        response = client.chat.completions.create(
            model=MODEL,
            messages=messages,  # type: ignore
            max_tokens=300,
            temperature=0.0
        )
        
        # 安全地获取响应内容
        result_text = ""
        if response and response.choices and len(response.choices) > 0:
            choice = response.choices[0]
            if choice and hasattr(choice, 'message'):
                message = choice.message
                if message and hasattr(message, 'content') and message.content is not None:
                    result_text = message.content.strip()
        
        logger.debug(f"AI 判断结果:\n{result_text}")

        # 解析结果：只要包含"不符合"就认为不通过
        if "不符合" in result_text:
            logger.debug("❌ 图片不符合探校主题")
            return False
        else:
            logger.debug("✅ 图片符合探校主题")
            return True

    except Exception as e:
        logger.error(f"AI 判断图片 {image_path} 时出错: {e}")
        # 出错时保守处理：保留图片（或可改为 False）
        return True  # 或 return False，根据你的策略


def download_reload_images(image_urls: List[str], text: str) -> None:
    """
    下载图片并在现有Word文档中追加内容

    参数：
    image_urls: 图片URL列表
    text: 文本内容（标题+内容）
    """
    logger.debug(f"开始下载 {len(image_urls)} 张图片")
    
    ## 确保images文件夹存在
    if not os.path.exists('images'):
        os.makedirs('images')
        logger.debug("创建images文件夹")

    ## 清空images文件夹
    logger.debug("清空images文件夹")
    for filename in os.listdir('images'):
        file_path = os.path.join('images', filename)
        try:
            if os.path.isfile(file_path):
                os.unlink(file_path)
                logger.debug(f"删除文件: {file_path}")
        except Exception as e:
            logger.error(f"删除文件 {file_path} 时出错: {e}")

    ## 下载图片并保存到images文件夹
    downloaded_images = []
    for i, url in enumerate(image_urls):
        try:
            logger.debug(f"下载图片 {i+1}/{len(image_urls)}: {url}")
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                # 生成文件名
                filename = f"image_{i+1}.jpg"
                file_path = os.path.join('images', filename)
                
                # 保存图片
                with open(file_path, 'wb') as f:
                    f.write(response.content)
                downloaded_images.append(file_path)
                logger.debug(f"已下载图片: {filename}")
            else:
                logger.error(f"下载图片失败，状态码: {response.status_code}")
        except Exception as e:
            logger.error(f"下载图片 {url} 时出错: {e}")

    ## 在现有Word文档中追加内容
    logger.debug("开始在现有Word文档中追加内容")
    try:
        # 打开现有的Word文档
        doc = Document('文案//教育//生成文案//finally_school_texts_images.docx')
        # 先筛选出符合“探校”主题的图片
        filtered_images = []
        for image_path in downloaded_images:
            time.sleep(2)
            if is_image_relevant_to_school(image_path):
                filtered_images.append(image_path)
            else:
                # 可选：删除不符合的图片以节省空间
                try:
                    os.remove(image_path)
                    logger.debug(f"已删除不符合主题的图片: {image_path}")
                except Exception as e:
                    logger.warning(f"删除图片 {image_path} 失败: {e}")

        # 插入筛选后的图片
        if filtered_images:
            for image_path in filtered_images:
                try:
                    if os.path.exists(image_path):
                        # 参考excel2word.py中的实现，添加更多错误处理
                        try:
                            # 尝试插入图片，如果失败则尝试不同的方式
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
                            run = paragraph.add_run()
                            # 使用较小的图片尺寸，避免内存问题
                            run.add_picture(image_path, width=Inches(4))
                            doc.add_paragraph()  # 添加空行
                            logger.debug(f"插入图片: {image_path}")
                        except Exception as insert_error:
                            # 如果直接插入失败，尝试使用BytesIO
                            try:
                                from PIL import Image
                                from io import BytesIO
                                # 打开图片并转换为RGB模式
                                with Image.open(image_path) as img:
                                    # 转换为RGB模式（如果是RGBA或其他模式）
                                    if img.mode in ('RGBA', 'LA', 'P'):
                                        img = img.convert('RGB')
                                    # 调整图片大小以减少内存使用
                                    max_size = (1200, 1200)  # 最大尺寸
                                    img.thumbnail(max_size, Image.Resampling.LANCZOS)
                                    # 保存为JPEG格式到内存
                                    img_buffer = BytesIO()
                                    img.save(img_buffer, format='JPEG', quality=85)
                                    img_buffer.seek(0)
                                    # 插入图片
                                    paragraph = doc.add_paragraph()
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run()
                                    run.add_picture(img_buffer, width=Inches(4))
                                    doc.add_paragraph()  # 添加空行
                                    logger.debug(f"通过PIL插入图片: {image_path}")
                            except Exception as pil_error:
                                logger.error(f"使用PIL插入图片 {image_path} 时出错: {pil_error}")
                                # 如果PIL也失败，尝试插入错误信息
                                error_para = doc.add_paragraph()
                                error_para.add_run(f"[图片插入失败: {os.path.basename(image_path)}]").bold = True
                    else:
                        logger.warning(f"图片文件不存在: {image_path}")
                except Exception as e:
                    logger.error(f"插入图片 {image_path} 时出错: {e}")
        
        # 然后插入标题和内容
        if text:
            # 分割标题和内容
            lines = text.split('\n')
            if lines:
                # 第一行作为标题，使用Heading 2样式
                heading = doc.add_heading(lines[0], level=2)
                # 其余行作为内容
                content = '\n'.join(lines[1:]) if len(lines) > 1 else ""
                if content:
                    doc.add_paragraph(content)
        
        # 添加分页符（如果不是第一个内容块）
        doc.add_page_break()
        
        # 保存文档
        try:
            doc.save('文案//教育//生成文案//finally_school_texts_images.docx')
            logger.debug("已更新Word文档")
        except Exception as e:
            logger.error(f"保存Word文档时出错: {e}")
            
    except Exception as e:
        logger.error(f"更新Word文档时出错: {e}")


if __name__ == "__main__":
    # 遍历文案\教育\生成文案\finally_school_texts.docx
    logger.debug("开始处理文档")
    try:
        doc = Document('文案//教育//生成文案//finally_school_texts.docx')
        logger.debug("成功加载主文档")
        
        # 检查是否存在带图片的文档，如果不存在则创建
        output_file = '文案//教育//生成文案//finally_school_texts_images.docx'
        if not os.path.exists(output_file):
            logger.debug("创建新的Word文档: 文案//教育//生成文案//finally_school_texts_images.docx")
            output_doc = Document()
            # 设置页面边距
            from docx.shared import Cm
            sections_output = output_doc.sections
            for section in sections_output:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
            # 保存初始空文档
            output_doc.save(output_file)
        else:
            logger.debug("使用现有的Word文档: 文案//教育//生成文案//finally_school_texts_images.docx")
        
        ## 获取所有的段落与标题名称(text=标题+内容)
        # 修改为根据Heading 2样式来分割内容
        sections = []
        current_section = {}
        
        # 遍历所有段落，根据Heading 2样式分割
        for paragraph in doc.paragraphs:
            # 检查段落样式是否为Heading 2
            if paragraph.style and paragraph.style.name == 'Heading 2':
                # 如果已经有当前章节，保存它
                if current_section:
                    sections.append(current_section)
                # 开始新章节
                current_section = {
                    'title': paragraph.text,
                    'content': ''
                }
            else:
                # 如果是普通段落，添加到当前章节内容中
                if current_section and paragraph.text.strip():
                    if current_section['content']:
                        current_section['content'] += '\n' + paragraph.text
                    else:
                        current_section['content'] = paragraph.text
        
        # 添加最后一个章节
        if current_section:
            sections.append(current_section)
        
        logger.debug(f"找到 {len(sections)} 个章节")
        
        # 检查已处理的章节数量，实现断点续跑功能
        processed_sections = 0
        if os.path.exists(output_file):
            try:
                output_doc_check = Document(output_file)
                # 统计Heading 2的数量来确定已处理的章节数
                heading_count = 0
                for paragraph in output_doc_check.paragraphs:
                    if paragraph.style and paragraph.style.name == 'Heading 2':
                        heading_count += 1
                processed_sections = heading_count
                logger.debug(f"已处理 {processed_sections} 个章节")
            except Exception as e:
                logger.warning(f"检查已处理章节时出错: {e}")
                processed_sections = 0
        
        # 处理每个章节（从已处理的章节开始）
        for i, section in enumerate(sections[processed_sections:], processed_sections):
            title = section['title']
            content = section['content']
            text = title + '\n' + content
            print("="*150)
            logger.debug(f"处理第 {i+1}/{len(sections)} 章节: {text[:50]}...")

            print("="*100)
            # 通过key_school_finded找到学校名称
            school_name = key_school_finded(text)
            if school_name:
                logger.debug(f"找到学校: {school_name}")

                print("="*100)
                # 通过get_school_image获取学校图片urls
                print("正在获取学校图片...")
                logger.debug("正在获取学校图片...")
                image_urls = get_school_image(school_name)
                logger.debug(f"获取到 {len(image_urls)} 张图片")

                print("="*100)
                # 通过download_reload_images下载并保存图片
                if image_urls:
                    download_reload_images(image_urls, text=text)
                else:
                    print("未获取到图片")
                    logger.debug("未获取到图片")
            else:
                print("未找到学校名称")
                logger.debug("未找到学校名称")
            
            time.sleep(15)
                
    except Exception as e:
        logger.error(f"处理文档时出错: {e}")
        print(f"处理文档时出错: {e}")
